"""
DOCX Parser Implementation

Uses python-docx for Word document text extraction.
"""
from typing import List, Dict, Any
from ...core.interfaces.document_parser import DocumentParserInterface


class DocxParser(DocumentParserInterface):
    """DOCX 文档解析器"""
    
    @property
    def supported_extensions(self) -> List[str]:
        return ["docx", "doc"]
    
    def parse(self, file_path: str) -> str:
        """
        解析 DOCX 文档，返回纯文本
        
        Args:
            file_path: DOCX 文件路径
            
        Returns:
            提取的纯文本内容
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except Exception as e:
            raise RuntimeError(f"DOCX parsing failed: {str(e)}")
    
    def parse_with_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        解析 DOCX 文档，返回文本和元数据
        
        Args:
            file_path: DOCX 文件路径
            
        Returns:
            包含 content 和 metadata 的字典
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # 提取文本
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # 提取元数据
            metadata = {}
            core_props = doc.core_properties
            if core_props:
                metadata = {
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else "",
                }
            metadata["paragraph_count"] = len(paragraphs)
            
            return {
                "content": "\n\n".join(paragraphs),
                "metadata": metadata
            }
        except Exception as e:
            raise RuntimeError(f"DOCX parsing failed: {str(e)}")
